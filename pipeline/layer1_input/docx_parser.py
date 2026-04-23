"""
DOCX 文档解析器
从Word文档提取文本、图片和元数据
"""

import os
import tempfile
from pathlib import Path
from typing import Optional

from models import RawContent, ImageData
from models.slide_spec import StructuredSection, TableData


class DocxParser:
    """DOCX文件解析器，使用python-docx"""

    def parse(self, file_path: str) -> RawContent:
        """
        解析DOCX文件为RawContent

        Args:
            file_path: .docx文件路径

        Returns:
            RawContent with source_type="doc"
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "缺少python-docx依赖，请运行: pip install python-docx"
            )

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"无法解析DOCX文件，文件可能已损坏: {e}")

        # 提取文本（任何子步骤失败都不应阻塞解析）
        try:
            raw_text = self._extract_text(doc)
        except Exception as e:
            print(f"⚠️  DOCX 文本提取部分失败，使用降级模式: {e}")
            raw_text = self._extract_text_fallback(doc)

        # 提取图片（图片不是必需的，任何失败都软降级为空列表）
        try:
            images = self._extract_images(doc, file_path)
        except Exception as e:
            print(f"⚠️  DOCX 图片提取失败，已跳过: {e}")
            images = []

        # 提取元数据
        try:
            metadata = self._extract_metadata(doc, file_path)
        except Exception:
            metadata = {"file_name": Path(file_path).name}

        # 检测语言和页面结构
        from .text_parser import TextParser
        tp = TextParser()
        lang = tp._detect_language(raw_text)
        source_pages, is_structured = tp.detect_structure(raw_text)

        # 结构化章节提取（利用段落样式）
        structured_sections = []
        try:
            structured_sections = self._extract_structure(doc)
        except Exception as e:
            print(f"[docx] 结构化提取失败，降级为扁平模式: {e}")

        # 如果text_parser没检测到结构，但docx有Heading段落，以docx结构为准
        if not is_structured and structured_sections:
            is_structured = True
            print(f"[docx] 从段落样式检测到{len(structured_sections)}个结构化章节")

        # 如果两种方式都没检测到足够结构，用文本内容做章节检测
        if not is_structured or len(source_pages) < 2:
            text_pages, text_structured = self._detect_text_sections(raw_text)
            if text_structured and len(text_pages) > len(source_pages):
                source_pages = text_pages
                if not is_structured:
                    is_structured = True
                    print(f"[docx] 从文本编号检测到{len(source_pages)}个结构化章节")

        # 补充：如果 structured_sections 太少但 source_pages 充足，从 source_pages 构建
        if len(structured_sections) < 2 and len(source_pages) >= 2:
            structured_sections = self._pages_to_sections(source_pages)
            print(f"[docx] 从source_pages构建{len(structured_sections)}个结构化章节")

        if is_structured:
            print(f"[docx] 结构化文档: {len(source_pages)}页, {len(structured_sections)}章节")

        return RawContent(
            source_type="doc",
            raw_text=raw_text,
            images=images,
            metadata=metadata,
            detected_language=lang,
            source_pages=source_pages,
            is_structured=is_structured,
            structured_sections=structured_sections,
        )

    def _extract_structure(self, doc) -> list[StructuredSection]:
        """
        利用python-docx段落样式提取层级章节树。
        Heading1/2/3 → 对应 level 1/2/3，其余段落归入最近的章节。
        表格归入其前方的最近章节。
        """
        # 标题样式名映射
        HEADING_MAP = {
            "Heading 1": 1, "Heading1": 1, "heading 1": 1,
            "Heading 2": 2, "Heading2": 2, "heading 2": 2,
            "Heading 3": 3, "Heading3": 3, "heading 3": 3,
            "Title": 0,  # 文档标题，不计入章节
        }

        # 第一遍：收集所有段落和标题位置
        elements = []  # (type, level_or_None, text)
        for para in doc.paragraphs:
            style_name = ""
            try:
                style_name = para.style.name if para.style else ""
            except Exception:
                pass

            text = (para.text or "").strip()
            if not text:
                continue

            if style_name in HEADING_MAP:
                elements.append(("heading", HEADING_MAP[style_name], text))
            else:
                elements.append(("body", None, text))

        # 表格也收集进来，标记在最后一个段落之后
        for table in doc.tables:
            try:
                headers = []
                rows = []
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        try:
                            cells.append((cell.text or "").strip())
                        except Exception:
                            cells.append("")
                    if not headers:
                        headers = cells
                    else:
                        rows.append(cells)
                if headers:
                    td = TableData(headers=headers, rows=rows)
                    elements.append(("table", None, td))
            except Exception:
                continue

        # 第二遍：构建章节树
        sections = []       # 顶级章节列表
        stack = []          # 当前章节栈 [(section, level), ...]

        for elem_type, level, content in elements:
            if elem_type == "heading" and level is not None and level > 0:
                section = StructuredSection(
                    title=content,
                    level=level,
                    content="",
                    char_count=0,
                )

                # 找到正确的父级
                while stack and stack[-1][1] >= level:
                    stack.pop()

                if stack:
                    stack[-1][0].children.append(section)
                else:
                    sections.append(section)

                stack.append((section, level))

            elif elem_type == "body":
                # 归入当前章节
                if stack:
                    current = stack[-1][0]
                    if current.content:
                        current.content += "\n\n" + content
                    else:
                        current.content = content
                    current.char_count = len(current.content)

            elif elem_type == "body":
                pass  # already handled above

            elif elem_type == "table":
                if stack:
                    stack[-1][0].tables.append(content)

        return sections if sections else []

    def _extract_text_fallback(self, doc) -> str:
        """文本提取降级：仅尝试段落，不碰表格"""
        lines = []
        try:
            for para in doc.paragraphs:
                t = (para.text or "").strip()
                if t:
                    lines.append(t)
        except Exception:
            pass
        return "\n\n".join(lines)

    def _extract_text(self, doc) -> str:
        """提取文档中的所有段落文本，对单个 cell/row 异常做容错"""
        paragraphs = []
        for para in doc.paragraphs:
            try:
                text = (para.text or "").strip()
            except Exception:
                continue
            if text:
                paragraphs.append(text)

        # 也提取表格中的文本（任何 cell 异常都跳过）
        for table in doc.tables:
            try:
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        try:
                            cells.append((cell.text or "").strip())
                        except Exception:
                            cells.append("")
                    line = " | ".join(cells)
                    if line.strip("| "):
                        paragraphs.append(line)
            except Exception:
                continue

        return "\n\n".join(paragraphs)

    # MIME → 文件扩展名
    _IMAGE_EXT_MAP = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/tiff": ".tiff",
        "image/bmp": ".bmp",
        "image/x-emf": ".emf",
        "image/x-wmf": ".wmf",
        "image/svg+xml": ".svg",
        "image/webp": ".webp",
    }

    def _extract_images(self, doc, file_path: str) -> list:
        """
        提取文档中的图片：document body + headers + footers。

        - 跳过外部链接（is_external=True）的图片关系，避免 ValueError
        - 单张图片解析失败不影响其它图片
        - 图片目录创建失败时，自动 fallback 到系统临时目录
        - 写入失败时跳过该张图但保留其它
        """
        images: list = []

        # 准备图片目录：优先文档同级 images/，失败则用系统 tmp
        img_dir = self._prepare_image_dir(file_path)
        if img_dir is None:
            print("⚠️  无法创建任何图片目录，跳过图片提取")
            return images

        # 收集所有需要遍历的 part：document, headers, footers, footnotes 等
        parts_to_scan = [doc.part]
        try:
            for section in doc.sections:
                for hf in (section.header, section.footer,
                           section.first_page_header, section.first_page_footer,
                           section.even_page_header, section.even_page_footer):
                    try:
                        if hf is not None and hasattr(hf, "part"):
                            parts_to_scan.append(hf.part)
                    except Exception:
                        continue
        except Exception:
            pass

        seen_targets = set()
        for part in parts_to_scan:
            try:
                rels = part.rels
            except Exception:
                continue
            for _rel_id, rel in rels.items():
                try:
                    if "image" not in rel.reltype:
                        continue
                    if getattr(rel, "is_external", False):
                        continue
                    target = rel.target_part
                    target_key = id(target)
                    if target_key in seen_targets:
                        continue
                    seen_targets.add(target_key)

                    image_blob = target.blob
                    content_type = getattr(target, "content_type", "") or ""
                    ext = self._IMAGE_EXT_MAP.get(content_type, ".png")

                    img_name = f"img_{len(images)}{ext}"
                    img_path = img_dir / img_name
                    try:
                        with open(img_path, "wb") as f:
                            f.write(image_blob)
                    except OSError as e:
                        print(f"⚠️  图片写入失败 {img_name}: {e}")
                        continue

                    images.append(ImageData(
                        file_path=str(img_path),
                        width_px=0,
                        height_px=0,
                        description="",
                    ))
                except Exception as e:
                    # 单张图片任意失败都跳过
                    print(f"⚠️  跳过单张图片 (rel={_rel_id}): {type(e).__name__}: {e}")
                    continue

        return images

    @staticmethod
    def _prepare_image_dir(file_path: str) -> Optional[Path]:
        """
        准备图片输出目录。优先文档同级 ./images/，写不动时回退系统 tmp。
        返回 None 表示两个目录都创建失败。
        """
        candidates = [
            Path(file_path).parent / "images",
            Path(tempfile.gettempdir()) / f"pptagent_images_{Path(file_path).stem}",
        ]
        for cand in candidates:
            try:
                cand.mkdir(parents=True, exist_ok=True)
                test_file = cand / ".write_test"
                test_file.write_bytes(b"x")
                test_file.unlink()
                return cand
            except Exception:
                continue
        return None

    def _detect_text_sections(self, raw_text: str) -> tuple[list, bool]:
        """用 TextParser 的编号检测能力，从文本内容中检测章节"""
        from .text_parser import TextParser
        tp = TextParser()
        pages = tp._detect_section_structure(raw_text)
        return pages, bool(pages)

    @staticmethod
    def _pages_to_sections(pages: list) -> list[StructuredSection]:
        """将 SourcePage 列表转为 StructuredSection 列表"""
        sections = []
        for p in pages:
            sections.append(StructuredSection(
                title=p.title,
                level=1,
                content=p.content,
                char_count=len(p.content),
                source_page_number=p.page_number,
            ))
        return sections

    def _extract_metadata(self, doc, file_path: str) -> dict:
        """提取文档元数据"""
        props = doc.core_properties
        meta = {
            "file_name": Path(file_path).name,
            "file_size": os.path.getsize(file_path),
            "paragraph_count": len(doc.paragraphs),
        }

        if props.title:
            meta["title"] = props.title
        if props.author:
            meta["author"] = props.author
        if props.created:
            meta["created"] = str(props.created)
        if props.modified:
            meta["modified"] = str(props.modified)

        return meta
